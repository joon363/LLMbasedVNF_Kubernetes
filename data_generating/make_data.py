from openai import OpenAI
from secret import OPENAI_API_KEY, KUBER_PASSWORD
from prompts import *
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

openai_client = OpenAI(api_key=OPENAI_API_KEY)
data_path='data/'

def make_ansible(success_count = 1):
    # Ansible 자동 생성. 일단 보류.
    for prompt in prompts['ansible']:
        response = openai_client.chat.completions.create(
            model = 'gpt-4o-mini', messages=[
                {"role" : "system", "content" : 'You are an expert in Kubernetes management.'},
                {"role" : "user", "content" : prompt}
                ], temperature=0)
        content=response.choices[0].message.content
        #print(content)
        #print('-------------------')
        yaml_pattern = r"```yaml(.*?)```"
        yaml_match = re.search(yaml_pattern, content, re.DOTALL)
        if yaml_match:
            yaml_content = yaml_match.group(1).strip()
            print("Extracted YAML content:")
            #print(yaml_content)

            yaml_file_path = data_path+"kubernetes_setup_"+str(success_count)+".yaml"
            with open(yaml_file_path, "w") as yaml_file:
                yaml_file.write(yaml_content)
            
            print(f"YAML content has been saved to {yaml_file_path}")
            
            # Ansible Playbook 실행
            import subprocess
            
            ansible_command = ["ansible-playbook", yaml_file_path, '--extra-vars', '"ansible_sudo_pass='+KUBER_PASSWORD+'"']
            
            try:
                result = subprocess.run(ansible_command, capture_output=True, text=True)
                if result.returncode == 0:
                    print("Ansible Playbook executed successfully.")
                    success_count += 1
                else:
                    print("Ansible Playbook execution failed.")
                    print("Error details:", result.stderr)
                    
                # 출력 결과
                print("Output:", result.stdout)
            except Exception as e:
                print("Failed to execute Ansible Playbook.")
                print(str(e))
        else:
            print("YAML content not found.")

    print(f'Total {success_count-1} YAML files have been created and executed successfully.')

def make_example_mop_prompt():
    example_mop_path='data/Example/'
    example_mop_list = os.listdir(example_mop_path)
    prompt='Here are the example MOP files.\n'
    for example_mop_file_name in example_mop_list:
        doc = Document(example_mop_path+example_mop_file_name)
        for para in doc.paragraphs:
            prompt+=para.text+'\n'
    return prompt

def make_language_prompt(lang):
    if lang=='ko':
        prompt='Please write in Korean'
    else:
        prompt='Please write in English'
    return prompt   

def make_function_prompt(function):
    prompt = ''
    if function=='firewall':
        prompt=' Use iptables operation, not ufw.'
    return prompt  

def make_system_prompt(system):
    prompt = ''
    if system=='OpenStack':
        prompt='Also, do not use the GUI(Horizon), use the CLI. '+ \
            'Instead of setting floating IP on the created VM, use the Jump Host, '+ \
            'which can connect to the internal VM, to connect to the newly created VM with SSH '+ \
            'and operate the shell commands in SSH connection. To do this, enable SSH access through the password. \n' + \
            "Don't make security groups or keypairs."
    return prompt

def make_formatted_prompt(function, additional_command):
    return = prompt.format( system=system_name, 
                            container=container_name, 
                            function=function, 
                            additional_command=additional_command)

def call_LLM(prompt):
    response = openai_client.chat.completions.create(
        model = 'gpt-4o-mini', messages=[
            {"role" : "system", "content" : f'You are an expert in {system_name} management.'},
            {"role" : "user", "content" : prompt}
            ], temperature=0)
    content=response.choices[0].message.content
    return content

def make_prompt(lang, function, additional_command, prompt):
    final_prompt = ''
    final_prompt += make_example_mop_prompt()
    final_prompt += make_formatted_prompt(function, additional_command)
    final_prompt += make_function_prompt(function)
    final_prompt += make_language_prompt(lang)
    final_prompt += make_system_prompt(system_name)
    return final_prompt

def make_mop_single(prompt):
    content = call_LLM(prompt)
    return content
    
def save_mop_single(mop, file_path):
    doc = Document()
    for line in mop.split('\n'):
        if line:
            if line.startswith('#'):
                head_num=line.count('#')
                doc.add_heading(line[head_num:], level=head_num)
            else:
                doc.add_paragraph(line)
    doc.save(file_path)

def make_mop():
    total_num_file=0
    
    for lang in ['en', 'ko']:
        for function in function_list:
            example_num=1
            for additional_command in additional_command_list[function]:
                for prompt in prompts['mop']:
                    # Setting prompt and path
                    final_prompt = make_prompt(lang, function, additional_command, prompt)
                    file_path = data_path+f"{system_name}_{function}_setup_{additional_command[0]}_{lang}_withExample_{example_num}.docx"
                    
                    # Make MOP, save into docx file
                    mop = make_mop_single(final_prompt)
                    save_mop_single(mop, file_path)

                    example_num+=1
                    total_num_file+=1

    print(f'Total {total_num_file} doc files have been created.')
